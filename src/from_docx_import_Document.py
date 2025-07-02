# src/from_docx_import_Document.py
from docx import Document
import os
import json

def read_first_text_in_file(file_path):
    """
    docxまたはtxtファイルを読み込み、最初に見つかったテキストを含む段落（または行）から100文字を出力する。
    """
    if not os.path.exists(file_path):
        print(f"エラー: ファイルが見つかりません - {file_path}")
        return
    try:
        if file_path.lower().endswith('.docx'):
            document = Document(file_path)
            for paragraph in document.paragraphs:
                # .strip()で空白や改行のみの段落を無視する
                if paragraph.text.strip():
                    print("最初に見つかったテキスト:")
                    print(paragraph.text[:100])
                    return paragraph.text
            print("ドキュメント内にテキストを含む段落が見つかりませんでした。")
            print("原因の可能性: 1. 全て空行である 2. テキストが段落ではなくテキストボックス等に含まれている")
        elif file_path.lower().endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                for line in f:
                    if line.strip():
                        print("最初に見つかったテキスト:")
                        print(line[:100])
                        return line.strip()
            print("テキストファイル内にテキストを含む行が見つかりませんでした。")
        else:
            print("対応していないファイル形式です。")
    except Exception as e:
        print(f"ファイルの読み込み中にエラーが発生しました: {e}")

def get_combined_knowledge_text(base_file_path: str, concepts_path: str) -> str:
    """
    base_file_path: ベースとなるdocxまたはtxtファイルのパス
    concepts_path: 構造化知識（concepts.jsonなど）のパス
    1. docx/txtの最初のテキスト
    2. 構造化知識の要約や要素
    を結合して返す
    """
    texts = []
    # 1. docx/txtの最初のテキスト
    if os.path.exists(base_file_path):
        try:
            if base_file_path.lower().endswith('.docx'):
                document = Document(base_file_path)
                for paragraph in document.paragraphs:
                    if paragraph.text.strip():
                        texts.append(paragraph.text.strip())
                        break
            elif base_file_path.lower().endswith('.txt'):
                with open(base_file_path, 'r', encoding='utf-8') as f:
                    for line in f:
                        if line.strip():
                            texts.append(line.strip())
                            break
            else:
                print("対応していないファイル形式です。")
        except Exception as e:
            print(f"ファイル読み込みエラー: {e}")
    # 2. 構造化知識の要約や要素
    if os.path.exists(concepts_path):
        try:
            with open(concepts_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            # concepts.jsonがリスト形式か単一オブジェクトか両対応
            if isinstance(data, dict):
                # 旧: conceptsキーあり
                if "concepts" in data and isinstance(data["concepts"], list):
                    for concept in data["concepts"]:
                        texts.append(concept.get("summary", ""))
                        if "components" in concept:
                            texts.extend([str(c) for c in concept["components"]])
                        texts.append(concept.get("implication", ""))
                # 新: conceptsキーなし
                else:
                    texts.append(data.get("summary", ""))
                    if "components" in data:
                        texts.extend([str(c) for c in data["components"]])
                    texts.append(data.get("implication", ""))
        except Exception as e:
            print(f"conceptsファイル読み込みエラー: {e}")
    return "\n".join([t for t in texts if t])

if __name__ == "__main__":
    # ご自身のファイルパスに修正してください
    file_path = "../data/knowledge_base/161217-master-Ryo.docx"  # または .txt も可
    read_first_text_in_file(file_path)
